"""Export approved project artifacts to markdown / json / text bundle / docx."""

from __future__ import annotations

import base64
import io
import json
from datetime import datetime, timezone
from uuid import UUID

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.domain.enums import ExportStatus, ExportType
from app.models import Artifact, Brief, ExportJob, Expert, ExpertFeedback, Project


async def create_export(
    db: AsyncSession, project_id: UUID, export_type: str
) -> ExportJob:
    job = ExportJob(
        project_id=project_id,
        export_type=export_type,
        status=ExportStatus.RUNNING.value,
    )
    db.add(job)
    await db.flush()

    try:
        project = await db.get(Project, project_id)
        if project is None:
            raise ValueError("Project not found")

        payload = await _collect(db, project)

        if export_type == ExportType.DOCX_SCENARIO.value:
            docx_payload = await _collect(db, project, approved_only=False)
            docx_bytes = _as_docx_scenario(docx_payload)
            content = {
                "docx_base64": base64.b64encode(docx_bytes).decode(),
                "filename": f"{project.title or 'scenario'}.docx",
            }
            path = f"exports/{project_id}/{job.id}.docx"
        elif export_type == ExportType.DOCX_PROFESSION_MAP.value:
            docx_payload = await _collect(db, project, approved_only=False)
            docx_bytes = _as_docx_profession_map(docx_payload)
            content = {
                "docx_base64": base64.b64encode(docx_bytes).decode(),
                "filename": f"{project.title or 'story'}_sujet.docx",
            }
            path = f"exports/{project_id}/{job.id}_story.docx"
        elif export_type == ExportType.JSON.value:
            content = payload
            path = f"exports/{project_id}/{job.id}.json"
        elif export_type == ExportType.TEXT_BUNDLE.value:
            content = {"files": _as_text_files(payload)}
            path = f"exports/{project_id}/{job.id}_bundle.json"
        else:
            content = {"markdown": _as_markdown(payload)}
            path = f"exports/{project_id}/{job.id}.md"

        job.result_content = content
        job.result_path = path
        job.status = ExportStatus.READY.value
    except Exception as exc:  # noqa: BLE001
        job.status = ExportStatus.FAILED.value
        job.error_message = str(exc)

    job.updated_at = datetime.now(timezone.utc)
    await db.flush()
    await db.refresh(job)
    return job


async def _collect(db: AsyncSession, project: Project, approved_only: bool = True) -> dict:
    brief = (
        await db.execute(select(Brief).where(Brief.project_id == project.id))
    ).scalar_one_or_none()
    q = select(Artifact).where(Artifact.project_id == project.id)
    if approved_only:
        q = q.where(Artifact.status == "approved")
    q = q.order_by(Artifact.step_type, Artifact.version.desc())
    artifacts = (await db.execute(q)).scalars().all()

    # Keep highest version per step
    latest: dict[str, Artifact] = {}
    for a in artifacts:
        if a.step_type not in latest:
            latest[a.step_type] = a

    experts = (
        await db.execute(select(Expert).where(Expert.project_id == project.id))
    ).scalars().all()
    feedback = (
        await db.execute(
            select(ExpertFeedback).where(ExpertFeedback.project_id == project.id)
        )
    ).scalars().all()

    preferred = ("profession_map", "scenario_plan")
    if any(k in latest for k in preferred):
        artifact_payload = {
            k: {"version": latest[k].version, "content": latest[k].content}
            for k in preferred
            if k in latest
        }
    else:
        artifact_payload = {
            k: {"version": v.version, "content": v.content} for k, v in latest.items()
        }

    return {
        "exported_at": datetime.now(timezone.utc).isoformat(),
        "project": {
            "id": str(project.id),
            "title": project.title,
            "client_name": project.client_name,
            "profession": project.profession,
            "audience": project.audience,
            "delivery_format": project.delivery_format,
            "expected_duration": project.expected_duration,
            "constraints": project.constraints,
        },
        "brief": brief.content_json if brief else {},
        "experts": [
            {"name": e.name, "role": e.role, "status": e.status} for e in experts
        ],
        "expert_feedback": [
            {"expert_id": str(f.expert_id), "content": f.content, "tags": f.structured_tags}
            for f in feedback
        ],
        "artifacts": artifact_payload,
    }


def _as_markdown(payload: dict) -> str:
    p = payload["project"]
    lines = [
        f"# {p['title']}",
        "",
        f"**Заказчик:** {p['client_name']}  ",
        f"**Профессия:** {p['profession']}  ",
        f"**Аудитория:** {p['audience']}  ",
        f"**Формат:** {p['delivery_format']}  ",
        f"**Длительность:** {p['expected_duration']}",
        "",
        "## Brief",
        "```json",
        json.dumps(payload.get("brief") or {}, ensure_ascii=False, indent=2),
        "```",
        "",
        "## Утверждённые артефакты",
    ]
    for step, art in (payload.get("artifacts") or {}).items():
        lines.append(f"### {step} (v{art['version']})")
        lines.append("```json")
        lines.append(json.dumps(art["content"], ensure_ascii=False, indent=2))
        lines.append("```")
        lines.append("")
    return "\n".join(lines)


def _as_text_files(payload: dict) -> dict[str, str]:
    files = {
        "project.json": json.dumps(payload["project"], ensure_ascii=False, indent=2),
        "brief.json": json.dumps(payload.get("brief") or {}, ensure_ascii=False, indent=2),
        "README.md": _as_markdown(payload),
    }
    for step, art in (payload.get("artifacts") or {}).items():
        files[f"{step}.json"] = json.dumps(art["content"], ensure_ascii=False, indent=2)
    return files


def _as_docx_scenario(payload: dict) -> bytes:
    """Generate a DOCX document for the scenario_plan artifact."""
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc = Document()

    # base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def _heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    def _kv(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value or "—")

    def _shade_row(row) -> None:
        """Apply light grey background to a header row."""
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "D9D9D9")
            tcPr.append(shd)

    def _make_table(headers: list[str]) -> object:
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        _shade_row(hdr)
        return tbl

    def _add_row(tbl, values: list[str]) -> None:
        row = tbl.add_row()
        for i, val in enumerate(values):
            row.cells[i].text = str(val or "")

    def _regs(scene: dict) -> None:
        regs = scene.get("regulations") or []
        if not regs:
            return
        reg_list = regs if isinstance(regs, list) else [regs]
        doc.add_paragraph().add_run("Регламенты:").bold = True
        for r in reg_list:
            doc.add_paragraph(str(r), style="List Bullet")

    def _props(scene: dict) -> None:
        val = scene.get("props")
        if not val:
            return
        if isinstance(val, list):
            val = ", ".join(str(x) for x in val)
        _kv("Реквизит", str(val))

    # ── cover page ────────────────────────────────────────────────────────
    proj = payload.get("project") or {}
    title_par = doc.add_heading(proj.get("title") or "Сценарий", 0)
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for label, key in [
        ("Заказчик", "client_name"),
        ("Профессия", "profession"),
        ("Аудитория", "audience"),
        ("Формат", "delivery_format"),
        ("Длительность", "expected_duration"),
    ]:
        if proj.get(key):
            _kv(label, proj[key])
    doc.add_page_break()

    # ── extract scenario sections ─────────────────────────────────────────
    artifacts = payload.get("artifacts") or {}
    content = (artifacts.get("scenario_plan") or {}).get("content") or {}
    sections = content.get("sections") or []

    training_items: list[dict] = []
    diagnostic_items: list[dict] = []
    for sec in sections:
        sid = sec.get("id", "")
        if sid == "training_scenes":
            training_items = sec.get("items") or []
        elif sid == "diagnostic_scenes":
            diagnostic_items = sec.get("items") or []

    # ── TRAINING SCENES ────────────────────────────────────────────────────
    if training_items:
        _heading("Обучающие сцены", level=1)
        for scene in training_items:
            _heading(scene.get("title") or "Сцена", level=2)
            if scene.get("actors"):
                _kv("Актёры", scene["actors"])
            if scene.get("location"):
                _kv("Локация", scene["location"])

            frames = scene.get("frames") or []
            if frames:
                doc.add_paragraph().add_run("Кадры:").bold = True
                tbl = _make_table(["№", "Действие в кадре", "Акцент"])
                for frame in frames:
                    _add_row(tbl, [
                        str(frame.get("shot_no", "")),
                        frame.get("action") or "",
                        frame.get("accent") or "",
                    ])
                doc.add_paragraph()

            if scene.get("audio_text"):
                p = doc.add_paragraph()
                p.add_run("Аудиотекст: ").bold = True
                p.add_run(scene["audio_text"])

            _regs(scene)
            _props(scene)
            doc.add_paragraph()

    # ── DIAGNOSTIC SCENES ──────────────────────────────────────────────────
    if diagnostic_items:
        if training_items:
            doc.add_page_break()
        _heading("Диагностические сцены", level=1)
        for scene in diagnostic_items:
            _heading(scene.get("title") or "Сцена", level=2)
            if scene.get("actors"):
                _kv("Актёры", scene["actors"])
            if scene.get("location"):
                _kv("Локация", scene["location"])

            frames = scene.get("frames") or []
            if frames:
                doc.add_paragraph().add_run("Кадры:").bold = True
                tbl = _make_table(["№", "Действие в кадре", "Нарушение (пусто = норма)", "Акцент"])
                for frame in frames:
                    _add_row(tbl, [
                        str(frame.get("shot_no", "")),
                        frame.get("action") or "",
                        frame.get("violation") or "",
                        frame.get("accent") or "",
                    ])
                doc.add_paragraph()

            vcats = scene.get("violation_categories") or []
            if vcats:
                doc.add_paragraph().add_run("Категории нарушений:").bold = True
                for cat in vcats:
                    if not isinstance(cat, dict):
                        continue
                    cat_par = doc.add_paragraph(style="List Bullet")
                    cat_par.add_run(cat.get("title") or "").bold = True
                    violation = cat.get("violation") or ""
                    if violation:
                        cat_par.add_run(f" — {violation}")
                    description = cat.get("description") or ""
                    if description:
                        doc.add_paragraph(description)

            _regs(scene)
            _props(scene)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _as_docx_profession_map(payload: dict) -> bytes:
    """Generate a DOCX document for the profession_map artifact."""
    from docx import Document  # type: ignore[import-untyped]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
    from docx.oxml import OxmlElement  # type: ignore[import-untyped]
    from docx.oxml.ns import qn  # type: ignore[import-untyped]
    from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def _heading(text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    def _kv(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value or "—")

    def _shade_row(row) -> None:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "D9D9D9")
            tcPr.append(shd)

    def _make_table(headers: list[str]) -> object:
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        _shade_row(hdr)
        return tbl

    def _add_row(tbl, values: list[str]) -> None:
        row = tbl.add_row()
        for i, val in enumerate(values):
            row.cells[i].text = str(val or "")

    proj = payload.get("project") or {}
    title_par = doc.add_heading(proj.get("title") or "Сюжет и точки оценки", 0)
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.add_run("Сюжет и точки оценки (промежуточная версия)").italic = True
    doc.add_paragraph()
    for label, key in [
        ("Заказчик", "client_name"),
        ("Профессия", "profession"),
        ("Аудитория", "audience"),
    ]:
        if proj.get(key):
            _kv(label, proj[key])
    doc.add_page_break()

    artifacts = payload.get("artifacts") or {}
    content = (artifacts.get("profession_map") or {}).get("content") or {}
    sections = content.get("sections") or []

    story_items: list[dict] = []
    assessment_items: list[dict] = []
    question_items: list[dict] = []
    for sec in sections:
        sid = sec.get("id", "")
        if sid == "work_storylines":
            story_items = sec.get("items") or []
        elif sid == "assessment_points":
            assessment_items = sec.get("items") or []
        elif sid == "expert_questions":
            question_items = sec.get("items") or []

    if story_items:
        _heading("Варианты работ и сюжет", level=1)
        for item in story_items:
            _heading(item.get("title") or "Вид работ", level=2)
            if item.get("description"):
                _kv("Описание", item["description"])
            steps = item.get("story_steps") or []
            if steps:
                doc.add_paragraph().add_run("Шаги сюжета:").bold = True
                for step in steps:
                    doc.add_paragraph(str(step), style="List Number")
            if item.get("attention_focus"):
                _kv("Фокус внимания", str(item["attention_focus"]))
            doc.add_paragraph()

    if assessment_items:
        doc.add_page_break()
        _heading("Навыки и точки оценки", level=1)
        for item in assessment_items:
            _heading(item.get("title") or "Вид работ", level=2)
            if item.get("description"):
                _kv("Контекст", item["description"])
            errors = item.get("errors") or []
            if errors:
                doc.add_paragraph().add_run("Ошибки:").bold = True
                tbl = _make_table(["Ошибка", "Правильно", "Визуальные признаки"])
                for err in errors:
                    if not isinstance(err, dict):
                        continue
                    cues = err.get("visual_cues") or []
                    if isinstance(cues, list):
                        cues_text = "; ".join(str(x) for x in cues)
                    else:
                        cues_text = str(cues)
                    _add_row(tbl, [err.get("error") or "", err.get("correct") or "", cues_text])
                doc.add_paragraph()

    if question_items:
        doc.add_page_break()
        _heading("Вопросы экспертам", level=1)
        for item in question_items:
            _heading(item.get("title") or "Вопрос", level=2)
            if item.get("description"):
                _kv("Контекст", item["description"])
            if item.get("why_needed"):
                _kv("Зачем нужен ответ", item["why_needed"])
            answer = item.get("answer")
            if answer:
                _kv("Ответ эксперта", str(answer))
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
